"""Document parsing: PDF/DOCX bytes in, per-page text out.

Page boundaries are preserved because citations are only useful if they point
at a page a person can turn to. PDFs have real pages. DOCX does not — Word
paginates at render time — so we split on explicit page breaks where the author
put them and fall back to a size-based split, and the README says so.

Scanned/image-only PDFs are out of scope for Phase 1: they parse to no text and
are rejected with a clear message rather than stored as an empty document.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass
from typing import List

import pdfplumber
from docx import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph

from .config import SUPPORTED_EXTENSIONS

# Roughly a page of dense contract prose. Only used for DOCX files that contain
# no explicit page breaks, so that "page 3" still means a consistent slice.
DOCX_CHARS_PER_PAGE = 2800


class IngestError(Exception):
    """Base class for problems the user can act on."""


class UnsupportedFileType(IngestError):
    pass


class EmptyDocumentError(IngestError):
    pass


class CorruptDocumentError(IngestError):
    pass


@dataclass
class PageText:
    page_number: int  # 1-based
    text: str


def parse_document(filename: str, data: bytes) -> List[PageText]:
    """Parse an uploaded file into pages of text.

    Raises UnsupportedFileType, CorruptDocumentError or EmptyDocumentError —
    all carry a message written for the end user.
    """
    extension = _extension_of(filename)
    if extension not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFileType(
            f"{filename}: only PDF and DOCX files are supported."
        )
    if not data:
        raise EmptyDocumentError(f"{filename}: the file is empty.")

    pages = _parse_pdf(filename, data) if extension == ".pdf" else _parse_docx(filename, data)

    pages = [PageText(p.page_number, _normalise(p.text)) for p in pages]
    if not any(p.text.strip() for p in pages):
        raise EmptyDocumentError(
            f"{filename}: no readable text was found. If this is a scanned "
            "document, it needs OCR, which this version does not support."
        )
    return pages


# ── PDF ──────────────────────────────────────────────────────────────────────


def _parse_pdf(filename: str, data: bytes) -> List[PageText]:
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            # Blank pages are kept so that page numbers stay aligned with the
            # numbering a reader sees in a PDF viewer.
            return [
                PageText(page_number=index, text=page.extract_text() or "")
                for index, page in enumerate(pdf.pages, start=1)
            ]
    except EmptyDocumentError:
        raise
    except Exception as exc:  # pdfplumber/pdfminer raise a wide range of types
        raise CorruptDocumentError(
            f"{filename}: the PDF could not be read. It may be corrupt or "
            "password-protected."
        ) from exc


# ── DOCX ─────────────────────────────────────────────────────────────────────


def _parse_docx(filename: str, data: bytes) -> List[PageText]:
    try:
        document = DocxDocument(io.BytesIO(data))
    except Exception as exc:
        raise CorruptDocumentError(
            f"{filename}: the Word document could not be read. It may be "
            "corrupt, or saved in the older .doc format."
        ) from exc

    # Blocks are the document's body in order, with an explicit-page-break flag.
    blocks: List[tuple[str, bool]] = []
    for element in _iter_block_items(document):
        if isinstance(element, Paragraph):
            text = element.text.strip()
            starts_new_page = _has_page_break(element)
            if text or starts_new_page:
                blocks.append((text, starts_new_page))
        elif isinstance(element, Table):
            rendered = _render_table(element)
            if rendered:
                blocks.append((rendered, False))

    if not blocks:
        return [PageText(page_number=1, text="")]

    has_explicit_breaks = any(flag for _, flag in blocks)
    return (
        _split_on_page_breaks(blocks)
        if has_explicit_breaks
        else _split_by_size(block for block, _ in blocks)
    )


def _split_on_page_breaks(blocks: List[tuple[str, bool]]) -> List[PageText]:
    pages: List[List[str]] = [[]]
    for text, starts_new_page in blocks:
        if starts_new_page and pages[-1]:
            pages.append([])
        if text:
            pages[-1].append(text)
    return [
        PageText(page_number=index, text="\n".join(lines))
        for index, lines in enumerate(pages, start=1)
    ]


def _split_by_size(blocks) -> List[PageText]:
    pages: List[List[str]] = [[]]
    length = 0
    for text in blocks:
        if length and length + len(text) > DOCX_CHARS_PER_PAGE:
            pages.append([])
            length = 0
        pages[-1].append(text)
        length += len(text) + 1
    return [
        PageText(page_number=index, text="\n".join(lines))
        for index, lines in enumerate(pages, start=1)
    ]


def _iter_block_items(document: DocxDocument):
    """Yield paragraphs and tables in document order.

    python-docx exposes .paragraphs and .tables as separate collections, which
    loses their relative position — walking the body XML keeps clauses and the
    tables that qualify them (schedules, price tiers) in the right order.
    """
    from docx.oxml.ns import qn

    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _has_page_break(paragraph: Paragraph) -> bool:
    """True if this paragraph is forced onto a new page."""
    from docx.oxml.ns import qn

    if paragraph.paragraph_format.page_break_before:
        return True
    for run in paragraph.runs:
        for br in run._element.findall(qn("w:br")):
            if br.get(qn("w:type")) == "page":
                return True
    return False


def _render_table(table: Table) -> str:
    """Flatten a table to pipe-separated rows so its values stay quotable."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        if any(cells):
            rows.append(" | ".join(cells))
    return "\n".join(rows)


# ── Shared ───────────────────────────────────────────────────────────────────


def _extension_of(filename: str) -> str:
    _, _, ext = filename.rpartition(".")
    return f".{ext.lower()}" if ext else ""


def _normalise(text: str) -> str:
    """Tidy whitespace without touching clause numbering or wording."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
